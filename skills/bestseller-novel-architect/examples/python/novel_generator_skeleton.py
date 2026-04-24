from pydantic import BaseModel, Field
from typing import List, Optional

# --- 1. DATA MODELS (Pydantic) ---
# Used to enforce structured output from the LLM

class Character(BaseModel):
    name: str
    goal: str = Field(description="The clear external goal")
    flaw_or_trauma: str = Field(description="Internal conflict or past trauma")
    arc_type: str = Field(description="Positive, Negative, or Complex")

class OutlineAct(BaseModel):
    act_number: int
    description: str
    key_plot_points: List[str]

class NovelOutline(BaseModel):
    genre: str
    premise: str
    characters: List[Character]
    acts: List[OutlineAct]

# --- 2. CORE PIPELINE ---
class NovelGenerator:
    def __init__(self, llm_client):
        self.llm = llm_client
        self.outline: Optional[NovelOutline] = None
        self.chapters: List[str] = []

    def generate_characters(self, genre: str, premise: str) -> List[Character]:
        """
        Calls LLM to generate characters based on Pillar 3 (Characteristics).
        Enforces external goal, internal flaw, and relatable details.
        """
        pass

    def generate_outline(self) -> NovelOutline:
        """
        Calls LLM to generate the 3-Act Structure (Pillar 1).
        Requires Setup (25%), Confrontation (50%), Resolution (25%).
        """
        pass

    def generate_chapter(self, chapter_number: int) -> str:
        """
        Calls LLM to generate a single chapter.
        MUST inject instructions for Pillar 6 (Chapter Structure) and Pillar 7 (Hooks/Cliffhangers).
        The prompt must enforce that the chapter ends with a Hook.
        """
        pass

    def export_to_docx(self, filename: str = "Novel_Output.docx"):
        """
        Export the generated novel to a formatted DOCX file.
        Requires: pip install python-docx
        """
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
        except ImportError:
            print("Please install python-docx: pip install python-docx")
            return

        doc = Document()
        
        # Title Page
        title = doc.add_heading(self.outline.premise if self.outline else "Generated Novel", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        # Chapters
        for i, chapter_text in enumerate(self.chapters, 1):
            heading = doc.add_heading(f"Chapter {i}", level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add chapter content with basic formatting
            for paragraph in chapter_text.split('\n'):
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph.strip())
                    # Standard novel formatting: first line indent, 12pt Times New Roman
                    p.paragraph_format.first_line_indent = Pt(12)
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                    
            doc.add_page_break()
            
        doc.save(filename)
        print(f"Successfully exported to {filename}")

# --- 3. SYSTEM PROMPTS (Constants) ---
# Inject these into your LLM calls to enforce the 7 Pillars
CHARACTER_PROMPT = """
You are a Bestseller Novel Architect. Create characters using the Iceberg Theory.
Every main character MUST have:
1. A clear external goal.
2. A psychological flaw or trauma.
...
"""

CHAPTER_PROMPT = """
Write Chapter {chapter_num}. 
MANDATORY RULES:
1. The very first line MUST be the chapter title in this exact format: "บทที่ {chapter_num}: [ชื่อบท]"
2. Start in medias res.
3. The middle must contain micro-conflict.
4. The chapter MUST end with a Cliffhanger or Hook (Action, Revelation, Question, Emotional, Shift, Promise).
5. Do NOT use western punctuation marks like `!` or `?`. Use Thai phrasing to convey tone.
DO NOT tie up loose ends.
"""

def main():
    # 1. Initialize LLM
    # 2. generator = NovelGenerator(llm)
    # 3. generator.generate_characters(...)
    # 4. generator.generate_outline()
    # 5. for i in range(total_chapters): 
    #        chapter = generator.generate_chapter(i)
    #        generator.chapters.append(chapter)
    # 6. generator.export_to_docx("My_Bestseller.docx")
    pass

if __name__ == "__main__":
    main()
