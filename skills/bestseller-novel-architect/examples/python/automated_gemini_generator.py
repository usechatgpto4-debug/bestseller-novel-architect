import os
import sys
import json
import time
from pydantic import BaseModel, Field
from typing import List

# ---------------------------------------------------------
# DEPENDENCIES:
# pip install google-genai pydantic python-docx
#
# ENVIRONMENT VARIABLE REQUIRED:
# set GEMINI_API_KEY=your_api_key_here
# ---------------------------------------------------------

try:
    from google import genai
    from google.genai import types
except ImportError:
    print("Error: google-genai is not installed. Please run: pip install google-genai")
    sys.exit(1)

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError:
    print("Error: python-docx is not installed. Please run: pip install python-docx")
    sys.exit(1)


# --- 1. DATA MODELS (Pydantic) ---
class Character(BaseModel):
    name: str = Field(description="ชื่อตัวละคร")
    goal: str = Field(description="เป้าหมายภายนอกที่ชัดเจน")
    flaw_or_trauma: str = Field(description="ความขัดแย้งภายใน หรือ ปมบาดแผลในอดีต (ทฤษฎีภูเขาน้ำแข็ง)")
    arc_type: str = Field(description="พัฒนาการของตัวละคร (เชิงบวก, เชิงลบ, หรือซับซ้อน)")

class OutlineAct(BaseModel):
    act_number: int
    description: str
    key_plot_points: List[str]

class NovelOutline(BaseModel):
    genre: str
    premise: str
    characters: List[Character]
    acts: List[OutlineAct]
    total_chapters: int

class SceneOutline(BaseModel):
    scene_number: int = Field(description="ลำดับของฉาก")
    setting: str = Field(description="สถานที่และเวลาที่เกิดฉากนี้ขึ้น")
    action_or_conflict: str = Field(description="เหตุการณ์หลัก ความขัดแย้ง หรือเป้าหมายย่อยในฉากนี้")

class ChapterOutline(BaseModel):
    chapter_number: int
    title: str = Field(description="ชื่อบทที่น่าดึงดูดใจเป็นภาษาไทย")
    scenes: List[SceneOutline] = Field(description="รายการฉากในบทนี้ (ปกติ 3-4 ฉาก)")


# --- 2. SYSTEM PROMPTS (7 Pillars Enforced) ---
MASTER_PROMPT = """คุณคือนักเขียนนิยายระดับเบสต์เซลเลอร์และเจ้าของรางวัลโนเบล
หน้าที่ของคุณคือการเขียนนิยายโดยยึดหลัก 8 เสาหลักดังนี้:
1. โครงสร้าง 3 องก์ (Setup, Confrontation, Resolution)
2. มุมมองการเล่าเรื่องแบบลึกซึ้ง (Deep POV - บุคคลที่ 1 หรือบุคคลที่ 3 แบบจำกัดมุมมอง) เพื่อสร้างความลุ้นระทึกและความใกล้ชิด
3. ตัวละครต้องมีเป้าหมายภายนอกที่ชัดเจนและมีบาดแผลทางจิตใจหรือปมในอดีต (ทฤษฎีภูเขาน้ำแข็ง)
4. จังหวะการเล่าเรื่อง (Pacing) ต้องกระชับ มีเป้าหมายย่อยๆ ในช่วงกลางเรื่อง
5. พัฒนาการของตัวละคร (Character Arc) ชัดเจน (เชิงบวก, เชิงลบ, หรือซับซ้อน)
6. โครงสร้างบท: เปิดฉากท่ามกลางเหตุการณ์ (In medias res) และมีความขัดแย้งย่อยในแต่ละบท
7. จุดตัดจบดึงดูดใจ (Chapter Hooks): ทุกบท 'ต้อง' จบด้วยความค้างคาหรือปริศนา (แอ็กชัน, การเปิดเผยความลับ, คำถาม, หรืออารมณ์) ห้ามคลี่คลายทุกอย่างในตอนจบของบท
8. แสดงให้เห็น อย่าแค่บอกเล่า (Show, Don't Tell): เน้นการใช้ประสาทสัมผัสทั้ง 5 (รูป, รส, กลิ่น, เสียง, สัมผัส) ห้ามสรุปอารมณ์ตรงๆ แต่ให้บรรยายปฏิกิริยาทางร่างกายและสรีรวิทยา ใช้คำกริยาที่ทรงพลัง
"""

# --- 3. CORE GENERATOR CLASS ---
class AutomatedNovelGenerator:
    def __init__(self, genre: str, premise: str, total_chapters: int):
        self.genre = genre
        self.premise = premise
        self.total_chapters = total_chapters
        
        # Initialize Gemini Client
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            print("ข้อผิดพลาด: ไม่พบตัวแปรสภาพแวดล้อม GEMINI_API_KEY")
            sys.exit(1)
            
        self.client = genai.Client()
        self.model_id = 'gemini-3.1-pro-preview' # As requested by user, must use preview when standard pro is unavailable
        
        self.outline: NovelOutline = None
        self.prologue_content: str = ""
        self.chapters_content: List[str] = []

    def generate_outline(self):
        print(">> ขั้นตอนที่ 1: กำลังสร้างข้อมูลตัวละครและโครงเรื่อง 3 องก์...")
        
        prompt = f"""
        จากประเภทของนิยาย '{self.genre}' และพล็อตเรื่อง '{self.premise}'
        โปรดสร้างโครงเรื่องนิยายอย่างละเอียดสำหรับ {self.total_chapters} บท
        รวมถึงสร้างข้อมูลตัวละครอย่างละเอียดโดยเน้นที่ปมทางจิตใจและเป้าหมายภายนอก
        และวางโครงเรื่องตามโครงสร้าง 3 องก์ (3-Act Structure)
        """
        
        max_retries = 3
        for attempt in range(max_retries):
            try:
                response = self.client.models.generate_content(
                    model=self.model_id,
                    contents=[MASTER_PROMPT, prompt],
                    config=types.GenerateContentConfig(
                        response_mime_type="application/json",
                        response_schema=NovelOutline,
                        temperature=0.7,
                    ),
                )
                
                self.outline = NovelOutline.model_validate_json(response.text)
                print(">> สร้างโครงเรื่องสำเร็จ!")
                break
            except Exception as e:
                print(f"   [!] เกิดข้อผิดพลาดในการสร้างโครงเรื่อง (พยายามครั้งที่ {attempt+1}/{max_retries}): {e}")
                if attempt == max_retries - 1:
                    print("   [!] ไม่สามารถสร้างโครงเรื่องได้หลังจากพยายามหลายครั้ง")
                    sys.exit(1)
                time.sleep(2)

    def _generate_and_critique_text(self, generation_prompt: str, log_name: str) -> str:
        max_retries = 3
        generated_text = ""
        
        # 1. Generate Raw Text
        for attempt in range(max_retries):
            try:
                response = self.client.models.generate_content(
                    model=self.model_id,
                    contents=[MASTER_PROMPT, generation_prompt],
                    config=types.GenerateContentConfig(temperature=0.8),
                )
                generated_text = response.text.strip()
                if generated_text:
                    break
            except Exception as e:
                print(f"   [!] เกิดข้อผิดพลาดในการสร้าง {log_name} (พยายามครั้งที่ {attempt+1}/{max_retries}): {e}")
                time.sleep(2)
                
        if not generated_text:
            print(f"   [!] ไม่สามารถสร้าง {log_name} ได้ ขอยกเลิกการทำงาน")
            sys.exit(1)

        # 2. Auto-Critique Loop
        print(f"      -> กำลังวิจารณ์และขัดเกลา {log_name}...")
        critique_prompt = f"""
        คุณคือบรรณาธิการมืออาชีพที่เข้มงวด ตรวจสอบข้อความฉากต่อไปนี้ที่เป็นภาษาไทย
        
        กฎที่ต้องบังคับใช้:
        1. ห้ามใช้เครื่องหมายวรรคตอนของภาษาอังกฤษ (`!` หรือ `?`) โดยเด็ดขาด ต้องลบออกและแทนที่ด้วยการบรรยายความรู้สึกหรือน้ำเสียงในภาษาไทย
        2. แสดงให้เห็น อย่าแค่บอกเล่า (Show, Don't Tell): ปรับปรุงการสรุปอารมณ์ที่อ่อนแอให้เป็นการบรรยายทางกายภาพโดยใช้ประสาทสัมผัสทั้ง 5
        3. ตรวจสอบให้แน่ใจว่าลักษณะนิสัยของตัวละครสอดคล้องกับคัมภีร์เรื่อง (Story Bible): {self.outline.model_dump_json()}
        
        หากข้อความละเมิดกฎเหล่านี้ ให้เขียนแก้ใหม่เพื่อแก้ไขปัญหา
        ผลลัพธ์ที่ตอบกลับให้ส่งมา 'เฉพาะ' ข้อความฉากที่ขัดเกลาแล้วเท่านั้น ห้ามใส่คำอธิบายการวิจารณ์ของคุณ
        
        ข้อความที่ต้องตรวจสอบ:
        {generated_text}
        """
        
        for attempt in range(max_retries):
            try:
                response = self.client.models.generate_content(
                    model=self.model_id,
                    contents=[MASTER_PROMPT, critique_prompt],
                    config=types.GenerateContentConfig(temperature=0.7),
                )
                polished_text = response.text.strip()
                if polished_text:
                    return polished_text
            except Exception as e:
                print(f"   [!] เกิดข้อผิดพลาดในการขัดเกลา {log_name} (พยายามครั้งที่ {attempt+1}/{max_retries}): {e}")
                time.sleep(2)
                
        return generated_text # Fallback to unpolished if critique fails

    def generate_prologue(self):
        print(">> ขั้นตอนที่ 2: กำลังสร้างบทนำ (Prologue)...")
        cache_file = "prologue_cache.txt"
        if os.path.exists(cache_file):
            print("   [!] พบไฟล์แคชของบทนำ ข้ามการสร้างใหม่")
            with open(cache_file, "r", encoding="utf-8") as f:
                self.prologue_content = f.read()
            return

        prompt = f"""
        เขียนบทนำ (Prologue) ที่น่าติดตามตื่นเต้นสำหรับนิยายเรื่องนี้
        
        คัมภีร์เรื่อง (Story Bible - Novel Outline JSON): {self.outline.model_dump_json()}
        ประเภท: {self.outline.genre}
        พล็อตเรื่อง: {self.outline.premise}
        
        กฎบังคับ:
        1. เขียนข้อความบทนำทั้งหมดเป็นภาษาไทย
        2. ดึงดูดผู้อ่านในทันทีด้วยปริศนาที่กระตุ้นความสนใจ แอ็กชัน หรือจุดเชื่อมโยงทางอารมณ์ที่ลึกซึ้ง
        3. รักษากระชับและทรงพลัง
        4. ผลลัพธ์ที่ตอบกลับให้ส่งมา 'เฉพาะ' ข้อความบทนำดิบๆ เท่านั้น ห้ามครอบด้วยบล็อกโค้ด markdown
        """
        
        content = self._generate_and_critique_text(prompt, "Prologue")
        self.prologue_content = content
        with open(cache_file, "w", encoding="utf-8") as f:
            f.write(content)
        print("   [บทนำเสร็จสมบูรณ์และถูกบันทึกในแคช]")

    def generate_chapters(self):
        print(f">> ขั้นตอนที่ 3: กำลังเขียนทั้ง {self.total_chapters} บทตามลำดับ (ทีละฉาก)...")
        
        for chapter_num in range(1, self.total_chapters + 1):
            print(f"\n--- กำลังเขียนบทที่ {chapter_num} ---")
            
            cache_file = f"chapter_{chapter_num}_cache.txt"
            if os.path.exists(cache_file):
                print(f"   [!] พบไฟล์แคชของบทที่ {chapter_num} ข้ามการสร้างใหม่")
                with open(cache_file, "r", encoding="utf-8") as f:
                    self.chapters_content.append(f.read())
                continue

            # 1. Generate Chapter Outline (Scenes)
            print("   >> กำลังสร้างโครงร่างฉาก...")
            previous_summary = "ยังไม่มีเนื้อหาก่อนหน้า"
            if self.chapters_content:
                previous_summary = "บทก่อนหน้าจบลงด้วยบริบทดังนี้:\n" + \
                                   self.chapters_content[-1][-500:]

            outline_prompt = f"""
            สร้างโครงร่างสำหรับบทที่ {chapter_num} จากทั้งหมด {self.total_chapters} บท
            
            คัมภีร์เรื่อง (Story Bible - Novel Outline JSON): {self.outline.model_dump_json()}
            บริบทของบทก่อนหน้า: {previous_summary}
            
            สร้างโครงร่างสำหรับบทที่ {chapter_num} โดยให้มี 3-4 ฉาก (scenes)
            พร้อมระบุชื่อบทที่น่าดึงดูดใจเป็นภาษาไทย
            """
            
            max_retries = 3
            chapter_outline = None
            for attempt in range(max_retries):
                try:
                    response = self.client.models.generate_content(
                        model=self.model_id,
                        contents=[MASTER_PROMPT, outline_prompt],
                        config=types.GenerateContentConfig(
                            response_mime_type="application/json",
                            response_schema=ChapterOutline,
                            temperature=0.7,
                        ),
                    )
                    chapter_outline = ChapterOutline.model_validate_json(response.text)
                    print(f"   >> สร้างโครงร่างบทที่ {chapter_num} สำเร็จ: '{chapter_outline.title}' ({len(chapter_outline.scenes)} ฉาก)")
                    break
                except Exception as e:
                    print(f"   [!] เกิดข้อผิดพลาดในการสร้างโครงร่างบท (พยายามครั้งที่ {attempt+1}/{max_retries}): {e}")
                    time.sleep(2)
            
            if not chapter_outline:
                print("   [!] ไม่สามารถสร้างโครงร่างบทได้ ขอยกเลิกการทำงาน")
                sys.exit(1)
                
            # 2. Generate and Critique Scenes
            full_chapter_text = f"บทที่ {chapter_num}: {chapter_outline.title}\n\n"
            scene_context = ""
            
            for scene in chapter_outline.scenes:
                print(f"      -> กำลังเขียนฉากที่ {scene.scene_number}: {scene.action_or_conflict[:40]}...")
                
                scene_prompt = f"""
                เขียนฉากที่ {scene.scene_number} สำหรับบทที่ {chapter_num}
                
                คัมภีร์เรื่อง (Story Bible - Novel Outline JSON): {self.outline.model_dump_json()}
                ชื่อบท: {chapter_outline.title}
                ฉากหลังและเวลา (Setting): {scene.setting}
                แอ็กชัน/ความขัดแย้งของฉาก: {scene.action_or_conflict}
                
                บริบทของฉากก่อนหน้า: {scene_context}
                
                กฎบังคับ:
                1. เขียนข้อความฉากทั้งหมดเป็นภาษาไทย
                2. แสดงให้เห็น อย่าแค่บอกเล่า (Show, Don't Tell): เน้นที่รายละเอียดทางประสาทสัมผัส
                3. ห้ามใช้เครื่องหมายวรรคตอนของภาษาอังกฤษ เช่น `!` หรือ `?` ให้ใช้การบรรยายอารมณ์แบบภาษาไทยแทน
                4. ผลลัพธ์ที่ตอบกลับให้ส่งมา 'เฉพาะ' ข้อความฉากดิบๆ เท่านั้น ห้ามครอบด้วยบล็อกโค้ด markdown
                """
                
                scene_text = self._generate_and_critique_text(scene_prompt, f"Scene {scene.scene_number}")
                
                full_chapter_text += scene_text + "\n\n***\n\n"
                scene_context = scene_text[-500:] # Last 500 chars for context
                
            # Clean up the trailing ***
            full_chapter_text = full_chapter_text.strip().rstrip('*').strip()
                
            self.chapters_content.append(full_chapter_text)
            
            # Cache the chapter
            with open(cache_file, "w", encoding="utf-8") as f:
                f.write(full_chapter_text)
                
            print(f"   [บทที่ {chapter_num} เสร็จสมบูรณ์และถูกบันทึกในแคช]")

    def export_to_docx(self, filename: str = "Detective_Novel_10_Chapters.docx"):
        print(f">> ขั้นตอนที่ 4: กำลังส่งออกเป็นไฟล์ {filename}...")
        
        doc = Document()
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Title Page
        title = doc.add_heading("รอยแค้นในความทรงจำสีจาง (Fading Echoes)", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        # Table of Contents (สารบัญ)
        toc_heading = doc.add_heading("สารบัญ", level=1)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if self.prologue_content:
            doc.add_paragraph("บทนำ (Prologue)")
        for i, chapter_text in enumerate(self.chapters_content, 1):
            lines = chapter_text.strip().split('\n')
            title = lines[0].strip() if lines else f"บทที่ {i}"
            doc.add_paragraph(title)
        doc.add_page_break()

        # Prologue
        if self.prologue_content:
            heading = doc.add_heading("บทนำ (Prologue)", level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in self.prologue_content.split('\n'):
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph.strip())
                    p.paragraph_format.first_line_indent = Pt(18)
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
            doc.add_page_break()

        # Chapters
        for i, chapter_text in enumerate(self.chapters_content, 1):
            lines = chapter_text.strip().split('\n')
            title = lines[0].strip() if lines else f"บทที่ {i}"
            content_lines = lines[1:] if lines else []
            
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for paragraph in content_lines:
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph.strip())
                    p.paragraph_format.first_line_indent = Pt(18)
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        
            doc.add_page_break()
            
        doc.save(filename)
        print(">> รวบรวมและสร้างไฟล์หนังสือสำเร็จ! (เสร็จสิ้น)")

def main():
    GENRE = "สืบสวนสอบสวน / ระทึกขวัญจิตวิทยา (Psychological Detective Thriller)"
    PREMISE = "สารวัตรธนา อดีตนักสืบมือฉมังที่กำลังเผชิญกับโรคความจำเสื่อมระยะเริ่มต้น (Early-onset Alzheimer's) ต้องเข้ามาสืบคดีฆาตกรรมต่อเนื่องในคฤหาสน์ปิดตาย ทว่าฆาตกรจงใจทิ้งเบาะแสที่เชื่อมโยงกับคดีสุดท้ายที่เขาเคยทำพลาดเมื่อ 10 ปีก่อน เขาต้องแข่งกับเวลาเพื่อหาตัวคนร้าย ก่อนที่ความทรงจำของเขาจะเลือนหายไปทั้งหมด พร้อมทั้งก้าวข้ามความรู้สึกผิดที่กัดกินจิตใจในอดีต"
    TOTAL_CHAPTERS = 10  

    print("=========================================")
    print(" Bestseller Novel Architect (Gemini API)")
    print("=========================================")
    
    generator = AutomatedNovelGenerator(
        genre=GENRE,
        premise=PREMISE,
        total_chapters=TOTAL_CHAPTERS
    )
    
    generator.generate_outline()
    generator.generate_prologue()
    generator.generate_chapters()
    generator.export_to_docx("Detective_Novel_10_Chapters.docx")

if __name__ == "__main__":
    main()
